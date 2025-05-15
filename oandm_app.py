from flask import Blueprint, request, jsonify, send_file
import os
import uuid
import traceback
from werkzeug.utils import secure_filename
from docx import Document
from openai import OpenAI

oandm_bp = Blueprint("oandm", __name__)
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

UPLOAD_FOLDER = "app/uploads"
GENERATED_FOLDER = "app/generated"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)

@oandm_bp.route("/ping", methods=["GET"])
def ping():
    return jsonify({"message": "O&M module is alive."})

@oandm_bp.route("/upload", methods=["POST"])
def upload_files():
    try:
        folder_id = str(uuid.uuid4())
        folder_path = os.path.join(UPLOAD_FOLDER, folder_id)
        os.makedirs(folder_path, exist_ok=True)

        files = request.files.getlist("files[]")
        if not files:
            return jsonify({"error": "No files received."}), 400

        for file in files:
            filename = secure_filename(file.filename)
            full_path = os.path.join(folder_path, filename)
            os.makedirs(os.path.dirname(full_path), exist_ok=True)
            file.save(full_path)

        # Start generating the document
        doc = Document()
        doc.add_heading("Operations & Maintenance Manual", 0)

        # Group files by first-level directory
        equipment_sections = {}
        for file in files:
            parts = file.filename.split("/")
            if len(parts) > 1:
                equipment = parts[0]
            else:
                equipment = "General"
            equipment_sections.setdefault(equipment, []).append(parts[-1])

        for equipment, file_list in equipment_sections.items():
            doc.add_page_break()
            doc.add_heading(f"Equipment: {equipment}", level=1)
            doc.add_paragraph("Included files:")
            for f in file_list:
                doc.add_paragraph(f"• {f}")

            # Generate GPT-based section
            prompt = f"""You are a professional process engineer creating an Operations & Maintenance manual.

Equipment: {equipment}
Files: {file_list}

Write a detailed and structured section for this equipment including:
- Purpose and description
- Installation overview
- Startup procedure
- Normal operation
- Shutdown procedure
- Maintenance intervals
- Spare parts or service notes

Be technical and realistic. Format with proper headings."""

            try:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "You are an expert in technical documentation for engineers."},
                        {"role": "user", "content": prompt}
                    ]
                )
                gpt_text = response.choices[0].message.content
                doc.add_paragraph(gpt_text)
            except Exception as gpt_error:
                doc.add_paragraph("⚠️ Failed to generate AI content.")
                print("GPT error for equipment:", equipment, gpt_error)

        # Save and return the document
        output_path = os.path.join(GENERATED_FOLDER, f"{folder_id}.docx")
        doc.save(output_path)
        return send_file(output_path, as_attachment=True)

    except Exception as e:
        print("🔥 Upload error:")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
